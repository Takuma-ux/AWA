from docx import Document
import os
import json
import argparse
import re

def save_image(image_part, output_folder, image_number):
    # 画像の拡張子を取得
    content_type = image_part.content_type
    ext = content_type.split('/')[-1]

    # 画像データを取得
    image_data = image_part.blob

    # 保存するファイル名
    file_name = f'image_{image_number}.{ext}'
    file_path = os.path.join(output_folder, file_name)

    # 画像を保存
    with open(file_path, 'wb') as f:
        f.write(image_data)
    print(f"画像を保存しました: {file_path}")

def remove_and_save_images_from_docx(file_path, output_folder):
    # ドキュメントを読み込む
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"ファイルを開く際にエラーが発生しました: {e}")
        return
    
    # 画像を保存するフォルダを作成
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    # 画像を抽出して保存
    image_number = 1
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            save_image(rel.target_part, output_folder, image_number)
            image_number += 1

    # 各段落内の画像を削除
    for para in doc.paragraphs:
        for run in para.runs:
            if 'Graphic' in run.element.xml:
                run.clear()

    # 各表内の画像を削除
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if 'Graphic' in run.element.xml:
                            run.clear()

    # 変更を保存
    output_file_path = file_path.replace('.docx', '_no_images.docx')
    try:
        doc.save(output_file_path)
        print(f"ファイルを保存しました: {output_file_path}")
    except Exception as e:
        print(f"ファイルを保存する際にエラーが発生しました: {e}")

    # 保存先のディレクトリとファイル名を確認
    print(f"保存先ディレクトリ: {os.path.dirname(output_file_path)}")
    print(f"保存ファイル名: {os.path.basename(output_file_path)}")

# コマンドライン引数をパースするための設定
parser = argparse.ArgumentParser(description='Process some files.')
parser.add_argument('--config', required=True, help='Path to the config JSON file')
args = parser.parse_args()

# JSONファイルのパスを取得
json_file_path = args.config
# JSONファイルで指定されている元のパスを取得
# JSONファイルを開いて読み込む
with open(json_file_path, 'r', encoding='utf-8-sig') as file:
    data = json.load(file)

base_dir = os.path.dirname(json_file_path) 
# 元のdocxファイルのパスを取得
docx_raw_file_path = os.path.abspath(os.path.join(base_dir, data["docx_raw_file_path"]))

# docx_file_path_2 を生成
docx_raw_file_name = os.path.basename(docx_raw_file_path)


# JSONファイル名から数字を抽出
match = re.search(r'\d+', os.path.basename(json_file_path))
if match:
    number = match.group()
else:
    number = 'default'  # 数字が見つからない場合のデフォルト値

output_dir = os.path.join(base_dir, "output", f"{number}")
if not os.path.exists(output_dir):
    os.makedirs(output_dir)
docx_file_name_modified = docx_raw_file_name.replace('.docx', '_without_toc_final.docx')
docx_file_path_2 = os.path.abspath(os.path.join(output_dir, docx_file_name_modified))
# output_folder の名前を設定
output_folder = os.path.join(output_dir, f'extracted_images')
remove_and_save_images_from_docx(docx_file_path_2, output_folder)


from docx import Document
import xml.etree.ElementTree as ET
import os
import colorsys  # HSV変換のために使用
import sys
import json
import argparse
import re

# 色範囲の定義
orange_rgb_range = {
    'r': (245, 260),  # Rの範囲
    'g': (210, 240),  # Gの範囲
    'b': (175, 225)   # Bの範囲
}

fill_blue_rgb_range = {
    'r': (178, 205),  # Rの範囲
    'g': (207, 220),  # Gの範囲
    'b': (234, 244)   # Bの範囲
}

def load_comments_from_file(file_path):
    comments_list = []

    # テキストファイルを読み込み
    with open(file_path, 'r', encoding='utf-8') as file:
        for line in file:
            # 行末の改行を削除し、「[」と「]」を取り除く
            line = line.strip().strip('[]')
            
            # 「コメント対象のテキスト」と「コメント内容」を分割
            comment_parts = line.split('],[')
            
            # 「」を取り除く
            cleaned_parts = [part.replace('「', '').replace('」', '') for part in comment_parts]
            
            # 二次元配列としてリストに格納
            comments_list.append(cleaned_parts)

    return comments_list

def load_headings(heading1_file, heading2_file):
    heading1_list = []
    heading2_list = []

    # 見出し1のファイルを読み込み
    with open(heading1_file, 'r', encoding='utf-8') as file:
        heading1_list = [line.strip() for line in file if line.strip()]  # 空行を除いてリストに追加

    # 見出し2のファイルを読み込み
    with open(heading2_file, 'r', encoding='utf-8') as file:
        heading2_group = []
        for line in file:
            line = line.strip()
            if line == "":  # 空行があった場合
                heading2_list.append(heading2_group)  # 現在のグループを追加
                heading2_group = []  # 次のグループの準備
            else:
                heading2_group.append(line)
        heading2_list.append(heading2_group)  # 最後のグループを追加

    # 見出し1に対応する見出し2を調整
    if len(heading1_list) > len(heading2_list):
        heading2_list.append([])  # 見出し1に対する見出し2がない場合に空のリストを追加

    return heading1_list, heading2_list

def hex_to_rgb(hex_color):
    """16進数カラーコードをRGBタプルに変換"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def rgb_to_hex(rgb_color):
    """RGBタプルを16進数カラーコードに変換"""
    return '#{:02x}{:02x}{:02x}'.format(rgb_color[0], rgb_color[1], rgb_color[2])

def is_rgb_in_range(rgb_color, color_range):
    """RGB値が指定された範囲内かどうかを確認"""
    return (color_range['r'][0] <= rgb_color[0] <= color_range['r'][1] and
            color_range['g'][0] <= rgb_color[1] <= color_range['g'][1] and
            color_range['b'][0] <= rgb_color[2] <= color_range['b'][1])

def rgb_to_hsv(rgb):
    """RGBからHSVに変換"""
    return colorsys.rgb_to_hsv(rgb[0]/255.0, rgb[1]/255.0, rgb[2]/255.0)

def is_blue_color(run):
    """テキストの色が青色かどうかを判定"""
    color_value = run.font.color.rgb
    if color_value:
        blue_rgb = (color_value[0], color_value[1], color_value[2])
        blue_hsv = rgb_to_hsv(blue_rgb)
        return blue_hsv[0] >= 200/360 and blue_hsv[0] <= 260/360  # 青色のHSV範囲（Hue）
    return False

def get_cell_background_color(cell):
    """セルの背景色を取得し、RGB範囲に基づいて適切な色に変換"""
    cell_xml = cell._tc
    tree = ET.ElementTree(ET.fromstring(cell_xml.xml))
    shd_elem = tree.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    
    if shd_elem is not None:
        fill_color = shd_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
        
        # fill_colorが6文字の16進数であることを確認
        if fill_color and len(fill_color) == 6 and all(c in '0123456789ABCDEFabcdef' for c in fill_color):
            rgb_color = hex_to_rgb(fill_color)
            
            if is_rgb_in_range(rgb_color, orange_rgb_range):
                return '#ffe4c4 '  # オレンジ系の範囲にある場合
            elif is_rgb_in_range(rgb_color, fill_blue_rgb_range):
                return '#C6D9F1'  # フィルブルー系の範囲にある場合

    return None  # 色が特定の範囲にない場合

def get_text_alignment_style(cell):
    """セル内のテキストのアライメントに応じてCSSスタイルを返す"""
    for paragraph in cell.paragraphs:
        alignment = paragraph.alignment
        print(f"Alignment: {alignment}")  # デバッグ用にアライメント値を出力
        if alignment == 1:  # 中央寄せ
            return "text-align: center;"
        elif alignment is None or alignment == 0 or alignment == 3:  # 左寄せまたは両端揃え
            return "text-align: left;"
    return ""  # デフォルトのアライメント（指定なし）

def create_html_tables(docx_path,links_file_path,heading1_file_path, heading2_file_path):
    doc = Document(docx_path)
    combined_html = ""
    comments_list = []
    comments_list = load_comments_from_file(links_file_path)
    heading1_array = []
    heading2_array = []
    heading1_array, heading2_array = load_headings(heading1_file_path, heading2_file_path)
    # 各表を処理する
    for table in doc.tables:
        # HTMLテーブルの開始
        html = '''<table style="table-layout: fixed; width: 100%; text-align: center; border-collapse: collapse;">
            <tbody>'''
        
        # 各行を処理する
        for i, row in enumerate(table.rows):
            html += '<tr>'
            for j, cell in enumerate(row.cells):
                # セルのテキストを取得
                cell_text = cell.text
                cell_text = cell_text.replace('\n','<br>\n')
                
                # セルの背景色を取得
                background_color = get_cell_background_color(cell)
                
                # セル内の最初のランの青色属性と太字属性をチェック
                has_blue_text = any(is_blue_color(run) for run in cell.paragraphs[0].runs)
                is_bold = any(run.bold for run in cell.paragraphs[0].runs)
                
                # テキストのアライメントを取得
                alignment_style = get_text_alignment_style(cell)
                
                # <th>タグか<td>タグを使用
                tag = "th" if i == 0 else "td"
                html += f'<{tag} style="background-color: {background_color or "#ffffff"}; {alignment_style}">'
                
                # 太字の場合は <strong> タグを使用
                if is_bold:
                    html += '<strong>'
                
                # 青色テキストの場合はリンクを追加
                if has_blue_text:
                    blue_text = ''
                    head1_flag = False
                    for comment in comments_list:
                        comment_text = comment[0]  # 「コメント対象のテキスト」を取得
                        if f"{comment_text}" in cell_text:
                            if "http" in comment[1]:
                                if "" in comment[1]:
                                    # link = comment[1].replace('','')
                                    blue_text = f'<a href="{comment[1]}">{comment[0]}</a>'
                                else:
                                    blue_text = f'<a href="{comment[1]}" target="_blank">{comment[0]}</a>'
                                # print(f"Found match for1: {comment_text}")
                            else:
                                for i, heading1 in enumerate(heading1_array):
                                    if heading1 in comment[1]:
                                        # print("見出し1へのリンクを発見")
                                        blue_text = f'<a href="#heading{i+1}">{comment[0]}</a>'
                                        head1_flag=True
                                        break
                                if not head1_flag:
                                    for row, sublist in enumerate(heading2_array):
                                        for col, heading2 in enumerate(sublist):
                                            if heading2 in comment[1]:
                                                # print("見出し2へのリンクを発見")
                                                blue_text = f'<a href="#heading{row+1}_{col+1}">{comment[0]}</a>'
                                # print(f"Found match for2: {comment_text}")
                            break
                    if not head1_flag:
                        blue_text = f'<a href="">{cell_text}</a>'
                    blue_text = blue_text.replace('<a','<span style="color: #56a0d6;"><a').replace('</a>','</a></span>')
                    html += blue_text
                else:
                    html += cell_text

                # 閉じタグ
                if is_bold:
                    html += '</strong>'
                
                html += f'</{tag}>'
            html += '</tr>'
        
        # HTMLテーブルの終了
        html += '''
            </tbody>
        </table>'''
        
        # テーブルのHTMLを追加し、改行を追加
        combined_html += html + '\n\n'
    
    return combined_html

# コマンドライン引数をパースするための設定
parser = argparse.ArgumentParser(description='Process some files.')
parser.add_argument('--config', required=True, help='Path to the config JSON file')
args = parser.parse_args()

# JSONファイルのパスを取得
json_file_path = args.config
# JSONファイルで指定されている元のパスを取得
# JSONファイルを開いて読み込む
with open(json_file_path, 'r', encoding='utf-8-sig') as file:
    data = json.load(file)

base_dir = os.path.dirname(json_file_path) 
# 元のdocxファイルのパスを取得
docx_raw_file_path = os.path.abspath(os.path.join(base_dir, data["docx_raw_file_path"]))
links_file_path = os.path.abspath(os.path.join(base_dir, data["links_file_path"]))
table_file_path = os.path.abspath(os.path.join(base_dir, data["table_file_path"]))
heading1_file_path = os.path.abspath(os.path.join(base_dir, data["heading1_file_path"]))
heading2_file_path = os.path.abspath(os.path.join(base_dir, data["heading2_file_path"]))
# docx_file_path_2 を生成
docx_raw_file_name = os.path.basename(docx_raw_file_path)
# JSONファイル名から数字を抽出
match = re.search(r'\d+', os.path.basename(json_file_path))
if match:
    number = match.group()
else:
    number = 'default'  # 数字が見つからない場合のデフォルト値

output_dir = os.path.join(base_dir, "output", f"{number}")
if not os.path.exists(output_dir):
    os.makedirs(output_dir)
docx_file_name_modified = docx_raw_file_name.replace('.docx', '_without_toc_final_no_images.docx')
docx_file_path_2 = os.path.abspath(os.path.join(output_dir, docx_file_name_modified))
# すべてのテーブルを1つのHTMLにまとめて保存
combined_html = create_html_tables(docx_file_path_2,links_file_path,heading1_file_path,heading2_file_path)
with open(table_file_path, "w", encoding="utf-8") as file:
    if combined_html:
        file.write(combined_html)
    else:
        file.write("")  # 空のファイルを作成
print(f"HTMLファイルが正常に作成されました: {table_file_path}")

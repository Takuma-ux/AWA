from docx import Document
from docx.shared import Pt
import os

def save_image(image_part, output_folder, image_number):
    """
    画像を指定されたフォルダに保存する関数。
    """
    content_type = image_part.content_type
    ext = content_type.split('/')[-1]
    image_data = image_part.blob

    file_name = f'image_{image_number}.{ext}'
    file_path = os.path.join(output_folder, file_name)

    with open(file_path, 'wb') as f:
        f.write(image_data)
    print(f"画像を保存しました: {file_path}")

def remove_images(doc):
    """
    ドキュメントからすべての画像を削除する関数。
    """
    # 各段落内の画像を削除
    for para in doc.paragraphs:
        for run in para.runs:
            if 'Graphic' in run.element.xml:
                run.clear()

def remove_until_first_title(doc):
    """
    Wordファイル内で最初のタイトル（太字かつ16pt）が出現するまでの段落およびテーブルを削除する関数。
    """
    delete_elements = []
    title_found = False

    # doc.element.body内のすべての要素を順に処理
    for element in list(doc.element.body):
        if not title_found:
            if element.tag.endswith('p'):  # 段落の場合
                for paragraph in doc.paragraphs:
                    if paragraph._element == element:
                        if is_title(paragraph):
                            title_found = True
                        else:
                            delete_elements.append(element)
                        break
            elif element.tag.endswith('tbl'):  # テーブルの場合
                delete_elements.append(element)

    # 収集した段落およびテーブルを削除
    for element in delete_elements:
        element.getparent().remove(element)

def is_title(paragraph):
    """
    段落が太字で、フォントサイズが16ptであるかを確認する関数。
    """
    for run in paragraph.runs:
        if run.bold and run.font.size == Pt(16):
            return True
    return False

def process_word_file(file_path, output_folder):
    """
    指定されたWordファイルを開いて、画像を保存し、その後最初のタイトルまでの内容を削除し、画像を削除して保存する関数。
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"ファイルを開く際にエラーが発生しました: {e}")
        return
    
    # 画像を保存する
    image_number = 1
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            save_image(rel.target_part, output_folder, image_number)
            image_number += 1

    # 最初のタイトルまでの段落やテーブルを削除
    remove_until_first_title(doc)
    
    # 画像を削除
    remove_images(doc)

    # 変更を保存
    output_file_path = file_path.replace('.docx', '_no_top_img.docx')
    try:
        doc.save(output_file_path)
        print(f"変更が保存されました: {output_file_path}")
    except Exception as e:
        print(f"ファイルを保存する際にエラーが発生しました: {e}")

# 使用例
file_path = r'C:\Users\nichi\OneDrive\デスクトップ\study\auto\240422【校了】登録販売者_役に立たない.docx'
output_folder = r'C:\Users\nichi\OneDrive\デスクトップ\study\auto\extracted_images'
process_word_file(file_path, output_folder)

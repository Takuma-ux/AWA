from docx import Document
import xml.etree.ElementTree as ET
import os

# 色範囲の定義
orange_rgb_range = {
    'r': (245, 255),  # Rの範囲
    'g': (220, 235),  # Gの範囲
    'b': (205, 225)   # Bの範囲
}

fill_blue_rgb_range = {
    'r': (178, 188),  # Rの範囲
    'g': (207, 217),  # Gの範囲
    'b': (234, 244)   # Bの範囲
}

def hex_to_rgb(hex_color):
    """16進数カラーコードをRGBタプルに変換"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def rgb_to_hex(rgb_color):
    """RGBタプルを16進数カラーコードに変換"""
    return '#{:02x}{:02x}{:02x}'.format(rgb_color[0], rgb_color[1], rgb_color[2])

def is_rgb_in_range(rgb_color, color_range):
    """RGB値が指定された範囲内かどうかを確認"""
    return (color_range['r'][0] <= rgb_color[0] <= color_range['r'][1] and
            color_range['g'][0] <= rgb_color[1] <= color_range['g'][1] and
            color_range['b'][0] <= rgb_color[2] <= color_range['b'][1])

def get_cell_background_color(cell):
    """セルの背景色を取得し、RGB範囲に基づいて適切な色に変換"""
    cell_xml = cell._tc
    tree = ET.ElementTree(ET.fromstring(cell_xml.xml))
    shd_elem = tree.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    
    if shd_elem is not None:
        fill_color = shd_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
        
        if fill_color:
            # 色が16進数形式（例: 'FFFF00'）と仮定
            rgb_color = hex_to_rgb(fill_color)
            
            if is_rgb_in_range(rgb_color, orange_rgb_range):
                return '#ffe8d1'  # オレンジ系の範囲にある場合
            elif is_rgb_in_range(rgb_color, fill_blue_rgb_range):
                return '#F0F8FF'  # フィルブルー系の範囲にある場合

    return None  # 色が特定の範囲にない場合

def create_html_table(docx_path, output_path):
    doc = Document(docx_path)
    
    # HTMLテーブルの開始
    html = '''<table style="table-layout: fixed; width: 100%; text-align: center; border-collapse: collapse;">
        <tbody>'''
    
    # 各行を処理する
    for i, row in enumerate(doc.tables[0].rows):
        html += '<tr>'
        for j, cell in enumerate(row.cells):
            # セルのテキストを取得
            cell_text = cell.text
            
            # セルの背景色を取得
            background_color = get_cell_background_color(cell)
            
            # セルの太字属性をチェック
            is_bold = any(run.bold for run in cell.paragraphs[0].runs)
            
            # <th>タグか<td>タグを使用
            if i == 0:
                # ヘッダー行の場合は<th>タグを使用
                if is_bold:
                    html += f'<th style="background-color: {background_color or "#ffe8d1"};"><strong>{cell_text}</strong></th>'
                else:
                    html += f'<th style="background-color: {background_color or "#ffe8d1"};">{cell_text}</th>'
            else:
                # データ行の場合は<td>タグを使用
                if j == 0:
                    # 1列目にはリンクを含むが、hrefは空
                    if is_bold:
                        html += f'<td style="background-color: {background_color or "#ffe8d1"};"><strong><a href="">{cell_text}</a></strong></td>'
                    else:
                        html += f'<td style="background-color: {background_color or "#ffe8d1"};"><a href="">{cell_text}</a></td>'
                else:
                    if is_bold:
                        html += f'<td style="background-color: {background_color or "#ffffff"};"><strong>{cell_text}</strong></td>'
                    else:
                        html += f'<td style="background-color: {background_color or "#ffffff"};">{cell_text}</td>'
        html += '</tr>'
    
    # HTMLテーブルの終了
    html += '''
        </tbody>
    </table>'''
    
    # HTMLファイルとして保存
    with open(output_path, "w", encoding="utf-8") as file:
        file.write(html)

    print("HTMLファイルが正常に作成されました。")

# 相対パスを設定
docx_path = os.path.join("input", "test_0802.docx")
output_path = os.path.join("output", "output.html")

create_html_table(docx_path, output_path)
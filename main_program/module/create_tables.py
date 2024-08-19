from docx import Document
import xml.etree.ElementTree as ET
import os
import colorsys  # HSV変換のために使用

# 色範囲の定義
orange_rgb_range = {
    'r': (245, 255),  # Rの範囲
    'g': (220, 235),  # Gの範囲
    'b': (205, 225)   # Bの範囲
}

fill_blue_rgb_range = {
    'r': (178, 188),  # Rの範囲
    'g': (207, 217),  # Gの範囲
    'b': (234, 244)   # Bの範囲
}

def hex_to_rgb(hex_color):
    """16進数カラーコードをRGBタプルに変換"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def rgb_to_hex(rgb_color):
    """RGBタプルを16進数カラーコードに変換"""
    return '#{:02x}{:02x}{:02x}'.format(rgb_color[0], rgb_color[1], rgb_color[2])

def is_rgb_in_range(rgb_color, color_range):
    """RGB値が指定された範囲内かどうかを確認"""
    return (color_range['r'][0] <= rgb_color[0] <= color_range['r'][1] and
            color_range['g'][0] <= rgb_color[1] <= color_range['g'][1] and
            color_range['b'][0] <= rgb_color[2] <= color_range['b'][1])

def rgb_to_hsv(rgb):
    """RGBからHSVに変換"""
    return colorsys.rgb_to_hsv(rgb[0]/255.0, rgb[1]/255.0, rgb[2]/255.0)

def is_blue_color(run):
    """テキストの色が青色かどうかを判定"""
    color_value = run.font.color.rgb
    if color_value:
        blue_rgb = (color_value[0], color_value[1], color_value[2])
        blue_hsv = rgb_to_hsv(blue_rgb)
        return blue_hsv[0] >= 200/360 and blue_hsv[0] <= 260/360  # 青色のHSV範囲（Hue）
    return False

def get_cell_background_color(cell):
    """セルの背景色を取得し、RGB範囲に基づいて適切な色に変換"""
    cell_xml = cell._tc
    tree = ET.ElementTree(ET.fromstring(cell_xml.xml))
    shd_elem = tree.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    
    if shd_elem is not None:
        fill_color = shd_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
        
        # fill_colorが6文字の16進数であることを確認
        if fill_color and len(fill_color) == 6 and all(c in '0123456789ABCDEFabcdef' for c in fill_color):
            rgb_color = hex_to_rgb(fill_color)
            
            if is_rgb_in_range(rgb_color, orange_rgb_range):
                return '#ffe8d1'  # オレンジ系の範囲にある場合
            elif is_rgb_in_range(rgb_color, fill_blue_rgb_range):
                return '#F0F8FF'  # フィルブルー系の範囲にある場合

    return None  # 色が特定の範囲にない場合

def create_html_tables(docx_path):
    doc = Document(docx_path)
    combined_html = ""
    
    # 各表を処理する
    for table in doc.tables:
        # HTMLテーブルの開始
        html = '''<table style="table-layout: fixed; width: 100%; text-align: center; border-collapse: collapse;">
            <tbody>'''
        
        # 各行を処理する
        for i, row in enumerate(table.rows):
            html += '<tr>'
            for j, cell in enumerate(row.cells):
                # セルのテキストを取得
                cell_text = cell.text

                cell_text = cell_text.replace('\n','<br>\n')
                
                # セルの背景色を取得
                background_color = get_cell_background_color(cell)
                
                # セル内の最初のランの青色属性と太字属性をチェック
                has_blue_text = any(is_blue_color(run) for run in cell.paragraphs[0].runs)
                is_bold = any(run.bold for run in cell.paragraphs[0].runs)
                
                # <th>タグか<td>タグを使用
                tag = "th" if i == 0 else "td"
                html += f'<{tag} style="background-color: {background_color or ("#ffe8d1" if i == 0 else "#ffffff")};">'
                
                # 太字の場合は <strong> タグを使用
                if is_bold:
                    html += '<strong>'
                
                # 青色テキストの場合はリンクを追加
                if has_blue_text:
                    html += f'<a href="">{cell_text}</a>'
                else:
                    html += cell_text

                # 閉じタグ
                if is_bold:
                    html += '</strong>'
                
                html += f'</{tag}>'
            html += '</tr>'
        
        # HTMLテーブルの終了
        html += '''
            </tbody>
        </table>'''
        
        # テーブルのHTMLを追加し、改行を追加
        combined_html += html + '\n\n'
    
    return combined_html

# 相対パスを設定
script_directory = os.path.dirname(os.path.abspath(__file__))
input_file_path = os.path.abspath(os.path.join(script_directory, '..','..', 'input', '240725_3.docx'))
output_file_path = os.path.abspath(os.path.join(script_directory, '..','..', 'output', 'combined_tables.html'))

# すべてのテーブルを1つのHTMLにまとめて保存
combined_html = create_html_tables(input_file_path)
with open(output_file_path, "w", encoding="utf-8") as file:
    file.write(combined_html)
print(f"HTMLファイルが正常に作成されました: {output_file_path}")

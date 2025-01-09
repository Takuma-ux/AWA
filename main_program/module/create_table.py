from docx import Document

# Wordファイルを読み込む
doc = Document("../input/test_0802.docx")

# HTMLテーブルの開始
html = '''<table style="table-layout: fixed; width: 100%; text-align: center; border-collapse: collapse;">
    <tbody>'''

# 各行を処理する
for i, row in enumerate(doc.tables[0].rows):
    html += '<tr>'
    for j, cell in enumerate(row.cells):
        # セルのテキストを取得
        cell_text = cell.text
        
        # セルの太字属性をチェック
        is_bold = any(run.bold for run in cell.paragraphs[0].runs)

        # <th>タグか<td>タグを使用
        if i == 0:
            # ヘッダー行の場合は<th>タグを使用
            if is_bold:
                html += f'<th style="background-color: #ffe8d1;"><strong>{cell_text}</strong></th>'
            else:
                html += f'<th style="background-color: #ffe8d1;">{cell_text}</th>'
        else:
            # データ行の場合は<td>タグを使用
            if j == 0:
                # 1列目にはリンクを含むが、hrefは空
                if is_bold:
                    html += f'<td style="background-color: #ffe8d1;"><strong><a href="">{cell_text}</a></strong></td>'
                else:
                    html += f'<td style="background-color: #ffe8d1;"><a href="">{cell_text}</a></td>'
            else:
                if is_bold:
                    html += f'<td><strong>{cell_text}</strong></td>'
                else:
                    html += f'<td>{cell_text}</td>'
    html += '</tr>'

# HTMLテーブルの終了
html += '''
    </tbody>
</table>'''

# HTMLファイルとして保存
with open("../output/output.html", "w", encoding="utf-8") as file:
    file.write(html)

print("HTMLファイルが正常に作成されました。")
